"""Document loading utilities supporting TXT, Markdown, PDF, and DOCX formats."""

from __future__ import annotations

import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """A loaded document with its content and metadata.

    Attributes:
        content: Raw text content of the document.
        source: Absolute path to the source file.
        metadata: Arbitrary key-value metadata (page numbers, author, etc.).
    """

    content: str
    source: str
    metadata: dict[str, object] = field(default_factory=dict)

    def __repr__(self) -> str:
        preview = self.content[:60].replace("\n", " ")
        return f"Document(source={Path(self.source).name!r}, preview={preview!r}...)"


class BaseLoader(ABC):
    """Abstract base class for document loaders."""

    @abstractmethod
    def load(self, path: Path) -> Document:
        """Load a single document from *path* and return a :class:`Document`."""

    def can_load(self, path: Path) -> bool:
        """Return *True* if this loader can handle the given file extension."""
        return path.suffix.lower() in self.supported_extensions

    @property
    @abstractmethod
    def supported_extensions(self) -> tuple[str, ...]:
        """File extensions handled by this loader (e.g. ``('.txt', '.md')``)."""


class PlainTextLoader(BaseLoader):
    """Loads ``.txt`` and ``.md`` files."""

    def __init__(self, encoding: str = "utf-8") -> None:
        self.encoding = encoding

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".txt", ".md")

    def load(self, path: Path) -> Document:
        text = path.read_text(encoding=self.encoding, errors="replace")
        return Document(
            content=text,
            source=str(path.resolve()),
            metadata={"extension": path.suffix, "size_bytes": path.stat().st_size},
        )


class PDFLoader(BaseLoader):
    """Loads ``.pdf`` files using *pypdf*."""

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".pdf",)

    def load(self, path: Path) -> Document:
        try:
            from pypdf import PdfReader  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError("Install pypdf: pip install pypdf") from exc

        reader = PdfReader(str(path))
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(text)
            logger.debug("Loaded PDF page %d from %s", i + 1, path.name)

        content = "\n\n".join(pages)
        return Document(
            content=content,
            source=str(path.resolve()),
            metadata={
                "extension": ".pdf",
                "num_pages": len(reader.pages),
                "size_bytes": path.stat().st_size,
            },
        )


class DocxLoader(BaseLoader):
    """Loads ``.docx`` files using *python-docx*."""

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)

    def load(self, path: Path) -> Document:
        try:
            from docx import Document as DocxDocument  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError("Install python-docx: pip install python-docx") from exc

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return Document(
            content=content,
            source=str(path.resolve()),
            metadata={"extension": ".docx", "num_paragraphs": len(paragraphs)},
        )


class DocumentLoader:
    """Dispatches to the appropriate :class:`BaseLoader` based on file extension.

    Example::

        loader = DocumentLoader()
        docs = list(loader.load_directory("data/sample_docs"))
    """

    _DEFAULT_LOADERS: tuple[BaseLoader, ...] = (
        PlainTextLoader(),
        PDFLoader(),
        DocxLoader(),
    )

    def __init__(self, loaders: tuple[BaseLoader, ...] | None = None) -> None:
        self._loaders: tuple[BaseLoader, ...] = loaders or self._DEFAULT_LOADERS

    def load(self, path: Path) -> Document:
        """Load a single file, raising :class:`ValueError` if unsupported."""
        path = Path(path)
        for loader in self._loaders:
            if loader.can_load(path):
                logger.info("Loading %s with %s", path.name, type(loader).__name__)
                return loader.load(path)
        supported = {ext for l in self._loaders for ext in l.supported_extensions}
        raise ValueError(
            f"No loader found for {path.suffix!r}. Supported: {sorted(supported)}"
        )

    def load_directory(
        self,
        directory: Path,
        *,
        recursive: bool = True,
    ) -> Iterator[Document]:
        """Yield :class:`Document` objects for every supported file in *directory*.

        Args:
            directory: Root directory to scan.
            recursive: If *True*, scan sub-directories as well.

        Yields:
            :class:`Document` instances in filesystem order.
        """
        directory = Path(directory)
        if not directory.is_dir():
            raise NotADirectoryError(f"{directory} is not a directory")

        pattern = "**/*" if recursive else "*"
        supported = {ext for l in self._loaders for ext in l.supported_extensions}

        for path in sorted(directory.glob(pattern)):
            if not path.is_file():
                continue
            if path.suffix.lower() not in supported:
                logger.debug("Skipping unsupported file %s", path.name)
                continue
            try:
                yield self.load(path)
            except Exception as exc:  # noqa: BLE001
                logger.warning("Failed to load %s: %s", path, exc)
